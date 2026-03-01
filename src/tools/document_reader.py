"""Universal document reader — text and image extraction for the document agent.

Provides dual-path extraction:
- Text path: for search, comparison (cheaper, faster)
- Vision path: for analysis, summarization (preserves layout)
"""

import asyncio
import csv
import hashlib
import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class Table:
    """Extracted table with headers and rows."""

    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    page: int | None = None

    def to_csv(self) -> str:
        out = io.StringIO()
        writer = csv.writer(out)
        if self.headers:
            writer.writerow(self.headers)
        writer.writerows(self.rows)
        return out.getvalue()

    def to_formatted_text(self) -> str:
        all_rows = ([self.headers] if self.headers else []) + self.rows
        if not all_rows:
            return ""
        col_widths = [
            max(len(str(row[i])) for row in all_rows if i < len(row))
            for i in range(max(len(r) for r in all_rows))
        ]
        lines = []
        for i, row in enumerate(all_rows):
            cells = [str(row[j]).ljust(col_widths[j]) for j in range(len(row))]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0 and self.headers:
                lines.append("|" + "|".join("-" * (w + 2) for w in col_widths) + "|")
        return "\n".join(lines)


async def extract_text(file_bytes: bytes, filename: str, mime_type: str) -> str:
    """Extract full text from any document. Routes by format."""
    ext = _get_extension(filename)

    if ext == "pdf":
        return await _extract_text_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return await _extract_text_docx(file_bytes)
    elif ext in ("xlsx", "xls"):
        return await _extract_text_xlsx(file_bytes)
    elif ext == "csv":
        return file_bytes.decode("utf-8", errors="replace")
    elif ext in ("txt", "md", "html", "rtf"):
        return file_bytes.decode("utf-8", errors="replace")
    elif mime_type and mime_type.startswith("image/"):
        return ""  # Images need OCR via Gemini, not text extraction
    else:
        return file_bytes.decode("utf-8", errors="replace")


async def extract_pages_as_images(
    file_bytes: bytes, filename: str, max_pages: int = 50
) -> list[bytes]:
    """Convert PDF pages to PNG images for Sonnet vision API.

    Args:
        max_pages: Maximum pages to render (default 50 to avoid OOM on huge PDFs).
    """
    ext = _get_extension(filename)
    if ext != "pdf":
        return []

    def _render():
        import pypdfium2

        images = []
        pdf = pypdfium2.PdfDocument(file_bytes)
        page_count = min(len(pdf), max_pages)
        for i in range(page_count):
            page = pdf[i]
            bitmap = page.render(scale=2)  # 2x for better OCR
            pil_image = bitmap.to_pil()
            buf = io.BytesIO()
            pil_image.save(buf, format="PNG")
            images.append(buf.getvalue())
            bitmap.close()
            page.close()
        pdf.close()
        return images

    return await asyncio.to_thread(_render)


async def extract_tables(file_bytes: bytes, filename: str, mime_type: str) -> list[Table]:
    """Extract tables from PDF/DOCX/XLSX/CSV."""
    ext = _get_extension(filename)

    if ext == "pdf":
        return await _extract_tables_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return await _extract_tables_docx(file_bytes)
    elif ext in ("xlsx", "xls"):
        return await _extract_tables_xlsx(file_bytes)
    elif ext == "csv":
        return [_parse_csv_table(file_bytes)]
    else:
        return []


async def get_page_count(file_bytes: bytes, filename: str) -> int:
    """Count pages in PDF/DOCX."""
    ext = _get_extension(filename)

    if ext == "pdf":

        def _count():
            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return len(pdf.pages)

        return await asyncio.to_thread(_count)
    elif ext in ("docx", "doc"):

        def _count_docx():
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            # Approximate: count section breaks + 1
            return max(1, len(doc.sections))

        return await asyncio.to_thread(_count_docx)
    return 1


async def extract_metadata(file_bytes: bytes, filename: str) -> dict:
    """Extract document metadata."""
    ext = _get_extension(filename)
    meta: dict = {
        "filename": filename,
        "file_size_bytes": len(file_bytes),
        "extension": ext,
    }

    if ext == "pdf":

        def _meta():
            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                meta["page_count"] = len(pdf.pages)
                if pdf.metadata:
                    for k in ("Title", "Author", "CreationDate", "Producer"):
                        if k in pdf.metadata and pdf.metadata[k]:
                            meta[k.lower()] = str(pdf.metadata[k])
            return meta

        return await asyncio.to_thread(_meta)

    return meta


async def compute_content_hash(file_bytes: bytes) -> str:
    """SHA-256 hash for duplicate detection."""
    return hashlib.sha256(file_bytes).hexdigest()


async def is_scanned_pdf(file_bytes: bytes) -> bool:
    """Check if a PDF is scanned (image-based) vs native text."""

    def _check():
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                total_chars = 0
                pages_checked = min(3, len(pdf.pages))
                for i in range(pages_checked):
                    text = pdf.pages[i].extract_text() or ""
                    total_chars += len(text.strip())
                return total_chars < 100 * pages_checked
        except Exception:
            return True  # If we can't parse, treat as scanned

    return await asyncio.to_thread(_check)


# --- Private helpers ---


def _get_extension(filename: str) -> str:
    if not filename:
        return ""
    parts = filename.rsplit(".", 1)
    return parts[-1].lower() if len(parts) > 1 else ""


async def _extract_text_pdf(file_bytes: bytes) -> str:
    def _extract():
        import pdfplumber

        texts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
        return "\n\n".join(texts)

    return await asyncio.to_thread(_extract)


async def _extract_text_docx(file_bytes: bytes) -> str:
    def _extract():
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    return await asyncio.to_thread(_extract)


async def _extract_text_xlsx(file_bytes: bytes) -> str:
    def _extract():
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"=== Sheet: {ws.title} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    parts.append("\t".join(cells))
        wb.close()
        return "\n".join(parts)

    return await asyncio.to_thread(_extract)


async def _extract_tables_pdf(file_bytes: bytes) -> list[Table]:
    def _extract():
        import pdfplumber

        tables = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                for raw_table in page.extract_tables():
                    if not raw_table or len(raw_table) < 2:
                        continue
                    headers = [str(c or "") for c in raw_table[0]]
                    rows = [[str(c or "") for c in row] for row in raw_table[1:]]
                    tables.append(Table(headers=headers, rows=rows, page=page_num + 1))
        return tables

    return await asyncio.to_thread(_extract)


async def _extract_tables_docx(file_bytes: bytes) -> list[Table]:
    def _extract():
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        tables = []
        for tbl in doc.tables:
            rows_data = []
            for row in tbl.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])
            if len(rows_data) >= 2:
                tables.append(Table(headers=rows_data[0], rows=rows_data[1:]))
            elif rows_data:
                tables.append(Table(rows=rows_data))
        return tables

    return await asyncio.to_thread(_extract)


async def _extract_tables_xlsx(file_bytes: bytes) -> list[Table]:
    def _extract():
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        tables = []
        for ws in wb.worksheets:
            rows_data = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows_data.append(cells)
            if len(rows_data) >= 2:
                tables.append(Table(headers=rows_data[0], rows=rows_data[1:]))
            elif rows_data:
                tables.append(Table(rows=rows_data))
        wb.close()
        return tables

    return await asyncio.to_thread(_extract)


def _parse_csv_table(file_bytes: bytes) -> Table:
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if len(rows) >= 2:
        return Table(headers=rows[0], rows=rows[1:])
    elif rows:
        return Table(rows=rows)
    return Table()
